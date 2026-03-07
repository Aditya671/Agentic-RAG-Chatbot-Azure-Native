# src/agentic_rag_chatbot_azure_native/indexer/llama_indexer.py

import os
import io
import hashlib
import uuid
import json
from datetime import datetime
from typing import List, Tuple, Dict, Optional, Iterable

import fitz  # PyMuPDF for PDFs
import pandas as pd
from docx import Document as DocxDocument  # python-docx
from llama_index import Document, GPTVectorStoreIndex, ServiceContext
from llama_index.embeddings import AzureOpenAIEmbedding
from llama_index.vector_stores import AzureAISearchVectorStore
from app_logger import setup_logger

logger, _ = setup_logger(name="llama-indexer")

# Config
CHUNK_SIZE = int(os.getenv("INDEX_CHUNK_SIZE", "1200"))  # characters
CHUNK_OVERLAP = int(os.getenv("INDEX_CHUNK_OVERLAP", "200"))
EMBED_BATCH_SIZE = int(os.getenv("EMBED_BATCH_SIZE", "16"))
INDEX_VERSION = "v1"

# Environment variables required:
# AZURE_OPENAI_API_KEY, AZURE_OPENAI_ENDPOINT, AZURE_OPENAI_EMBED_DEPLOYMENT
# AZURE_SEARCH_SERVICE_NAME, AZURE_SEARCH_API_KEY, AZURE_SEARCH_INDEX_NAME

def compute_checksum_bytes(data: bytes) -> str:
    h = hashlib.sha256()
    h.update(data)
    return h.hexdigest()

def compute_checksum_file(path: str) -> str:
    with open(path, "rb") as f:
        return compute_checksum_bytes(f.read())

def make_doc_id(source: str) -> str:
    # deterministic UUID5 based on source path/URL
    return str(uuid.uuid5(uuid.NAMESPACE_URL, source))

def now_iso() -> str:
    return datetime.utcnow().isoformat() + "Z"

def chunk_text_simple(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[Tuple[str, int, int]]:
    """
    Simple character-based chunking with overlap.
    Returns list of (chunk_text, start_offset, end_offset)
    """
    if not text:
        return []
    chunks = []
    start = 0
    text_len = len(text)
    step = chunk_size - overlap
    if step <= 0:
        step = chunk_size
    while start < text_len:
        end = min(start + chunk_size, text_len)
        chunk = text[start:end]
        chunks.append((chunk, start, end))
        if end == text_len:
            break
        start += step
    return chunks

# --- Extractors ---

def extract_text_from_pdf(path: str) -> Tuple[str, int]:
    doc = fitz.open(path)
    parts = []
    for page in doc:
        parts.append(page.get_text("text"))
    return "\n".join(parts), len(doc)

def extract_text_from_docx(path: str) -> Tuple[str, None]:
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    return "\n".join(paragraphs), None

def extract_text_from_txt(path: str) -> Tuple[str, None]:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read(), None

def extract_text_from_csv(path: str, text_columns: Optional[List[str]] = None, max_rows: Optional[int] = None) -> Tuple[str, None]:
    df = pd.read_csv(path)
    if text_columns:
        df = df[text_columns]
    if max_rows:
        df = df.head(max_rows)
    # Convert to a readable text representation
    return df.to_csv(index=False), None

def extract_text_from_dataframe(df: pd.DataFrame, text_columns: Optional[List[str]] = None, max_rows: Optional[int] = None) -> Tuple[str, None]:
    if text_columns:
        df = df[text_columns]
    if max_rows:
        df = df.head(max_rows)
    return df.to_csv(index=False), None

# --- Llama-index initialization ---

def init_embedding_and_vectorstore():
    deployment = os.environ.get("AZURE_OPENAI_EMBED_DEPLOYMENT", "text-embedding-3-large")
    api_key = os.environ.get("AZURE_OPENAI_API_KEY")
    endpoint = os.environ.get("AZURE_OPENAI_ENDPOINT")

    embedding = AzureOpenAIEmbedding(
        deployment=deployment,
        api_key=api_key,
        endpoint=endpoint,
    )

    service_name = os.environ.get("AZURE_SEARCH_SERVICE_NAME")
    service_key = os.environ.get("AZURE_SEARCH_API_KEY")
    index_name = os.environ.get("AZURE_SEARCH_INDEX_NAME")

    vector_store = AzureAISearchVectorStore(
        service_name=service_name,
        api_key=service_key,
        index_name=index_name,
        embedding=embedding,
    )

    service_context = ServiceContext.from_defaults(embed_model=embedding)
    return embedding, vector_store, service_context

# --- Document builder ---

def build_doc_metadata_for_file(path: str, checksum: str, extra: Optional[Dict] = None) -> Dict:
    stat = os.stat(path)
    meta = {
        "doc_id": make_doc_id(path),
        "source_path": path,
        "filename": os.path.basename(path),
        "checksum": checksum,
        "uploaded_date": now_iso(),
        "last_modified": datetime.utcfromtimestamp(stat.st_mtime).isoformat() + "Z",
        "file_size": stat.st_size,
        "mime_type": _guess_mime(path),
        "index_version": INDEX_VERSION,
        "ingest_pipeline": "llama-indexer-v1",
    }
    if extra:
        meta.update(extra)
    return meta

def _guess_mime(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return "application/pdf"
    if ext in (".txt", ".md"):
        return "text/plain"
    if ext in (".csv",):
        return "text/csv"
    if ext in (".docx",):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return "application/octet-stream"

def create_documents_from_text(text: str, base_meta: Dict) -> List[Document]:
    chunks = chunk_text_simple(text)
    docs: List[Document] = []
    for i, (chunk_text, start, end) in enumerate(chunks):
        chunk_id = f"{base_meta['doc_id']}::chunk::{i}"
        meta = dict(base_meta)
        meta.update({
            "chunk_id": chunk_id,
            "chunk_start_offset": start,
            "chunk_end_offset": end,
            "chunk_start_page": base_meta.get("chunk_start_page"),
            "chunk_end_page": base_meta.get("chunk_end_page"),
            "indexed_at": now_iso(),
        })
        docs.append(Document(text=chunk_text, metadata=meta))
    return docs

# --- High-level indexer functions ---

def index_file(path: str, force_reindex: bool = False, vector_store=None, service_context=None) -> Dict:
    """
    Index a single file. Returns metadata summary.
    """
    logger.info(f"Indexing file: {path}")
    checksum = compute_checksum_file(path)
    meta = build_doc_metadata_for_file(path, checksum)

    # TODO: consult metadata DB to skip if checksum unchanged unless force_reindex True
    # Example: if not force_reindex and metadata_db.exists(doc_id) and metadata_db[doc_id]['checksum'] == checksum: skip

    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        text, page_count = extract_text_from_pdf(path)
        meta["page_count"] = page_count
    elif ext == ".docx":
        text, _ = extract_text_from_docx(path)
    elif ext in (".txt", ".md"):
        text, _ = extract_text_from_txt(path)
    elif ext == ".csv":
        text, _ = extract_text_from_csv(path)
    else:
        # fallback: try reading as text
        try:
            text, _ = extract_text_from_txt(path)
        except Exception:
            raise ValueError(f"Unsupported file type: {ext}")

    docs = create_documents_from_text(text, meta)

    # init vector store if not provided
    if vector_store is None or service_context is None:
        _, vector_store, service_context = init_embedding_and_vectorstore()

    # Build index and upsert
    logger.info(f"Upserting {len(docs)} chunks for doc_id {meta['doc_id']}")
    index = GPTVectorStoreIndex.from_documents(docs, service_context=service_context, vector_store=vector_store)
    # Optionally persist index to disk or metadata DB
    logger.info("Upsert complete")
    return {"doc_id": meta["doc_id"], "chunks_indexed": len(docs), "checksum": checksum}

def index_dataframe(df: pd.DataFrame, name: str = "dataframe", text_columns: Optional[List[str]] = None, vector_store=None, service_context=None) -> Dict:
    """
    Index a pandas DataFrame. Converts to CSV text representation by default.
    """
    logger.info("Indexing DataFrame")
    text, _ = extract_text_from_dataframe(df, text_columns)
    checksum = compute_checksum_bytes(text.encode("utf-8"))
    meta = {
        "doc_id": make_doc_id(name + checksum),
        "source_path": name,
        "filename": name,
        "checksum": checksum,
        "uploaded_date": now_iso(),
        "last_modified": now_iso(),
        "file_size": len(text),
        "mime_type": "text/csv",
        "index_version": INDEX_VERSION,
        "ingest_pipeline": "dataframe-csv",
    }
    docs = create_documents_from_text(text, meta)
    if vector_store is None or service_context is None:
        _, vector_store, service_context = init_embedding_and_vectorstore()
    index = GPTVectorStoreIndex.from_documents(docs, service_context=service_context, vector_store=vector_store)
    logger.info("DataFrame indexing complete")
    return {"doc_id": meta["doc_id"], "chunks_indexed": len(docs), "checksum": checksum}

# --- Search helper ---

def semantic_search(query: str, top_k: int = 5, vector_store=None, service_context=None):
    """
    Simple semantic search wrapper using the vector store index.
    """
    if vector_store is None or service_context is None:
        _, vector_store, service_context = init_embedding_and_vectorstore()
    # Create a temporary index for querying
    index = GPTVectorStoreIndex.from_documents([], service_context=service_context, vector_store=vector_store)
    response = index.as_query_engine().query(query)
    return response

# --- CLI helper ---

def index_path(path: str, recursive: bool = True, force: bool = False):
    """
    Walk a path and index supported files.
    """
    _, vector_store, service_context = init_embedding_and_vectorstore()
    supported = {".pdf", ".docx", ".txt", ".md", ".csv"}
    results = []
    if os.path.isfile(path):
        results.append(index_file(path, force_reindex=force, vector_store=vector_store, service_context=service_context))
    else:
        for root, dirs, files in os.walk(path):
            for f in files:
                ext = os.path.splitext(f)[1].lower()
                if ext in supported:
                    p = os.path.join(root, f)
                    try:
                        results.append(index_file(p, force_reindex=force, vector_store=vector_store, service_context=service_context))
                    except Exception as e:
                        logger.error(f"Failed to index {p}: {e}")
            if not recursive:
                break
    return results